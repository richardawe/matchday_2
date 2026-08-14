from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
from pathlib import Path

OUT = Path(__file__).resolve().parent / "Matchday_Africa_App_Feature_and_Operations_Guide.docx"
NAVY=RGBColor(25,30,27); GOLD=RGBColor(174,127,59); MUTED=RGBColor(92,96,89); WHITE=RGBColor(255,255,255); PALE=RGBColor(244,240,231)

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def set_table_geometry(table, widths):
    table.autofit=False; tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells): set_cell_width(cell,widths[i]); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_font(run,name='Aptos',size=10.5,bold=None,color=NAVY,italic=None):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size); run.font.color.rgb=color
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic

def page_number(paragraph):
    run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.82); sec.bottom_margin=Inches(.78); sec.left_margin=Inches(.82); sec.right_margin=Inches(.82); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); normal.font.size=Pt(10.5); normal.font.color.rgb=NAVY; normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.18
for name,size,before,after,color in [('Title',30,0,10,NAVY),('Subtitle',13,0,18,MUTED),('Heading 1',18,18,8,GOLD),('Heading 2',14,13,6,NAVY),('Heading 3',11.5,9,4,GOLD)]:
    s=styles[name]; s.font.name='Aptos Display' if name!='Subtitle' else 'Aptos'; s._element.rPr.rFonts.set(qn('w:ascii'),s.font.name); s._element.rPr.rFonts.set(qn('w:hAnsi'),s.font.name); s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=name!='Subtitle'; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for lname in ['List Bullet','List Number']:
    styles[lname].font.name='Aptos'; styles[lname].font.size=Pt(10.5); styles[lname].paragraph_format.left_indent=Inches(.38); styles[lname].paragraph_format.first_line_indent=Inches(-.19); styles[lname].paragraph_format.space_after=Pt(4); styles[lname].paragraph_format.line_spacing=1.18
if 'Callout' not in styles:
    cs=styles.add_style('Callout',WD_STYLE_TYPE.PARAGRAPH); cs.base_style=normal; cs.font.name='Aptos'; cs.font.size=Pt(11); cs.font.color.rgb=NAVY; cs.paragraph_format.left_indent=Inches(.2); cs.paragraph_format.right_indent=Inches(.2); cs.paragraph_format.space_before=Pt(8); cs.paragraph_format.space_after=Pt(8)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=header.add_run('MATCHDAY AFRICA  /  PRODUCT & OPERATIONS GUIDE'); set_font(r,size=8,bold=True,color=MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run('Internal product reference  ·  '); set_font(r,size=8,color=MUTED); page_number(footer)

def p(text='',style=None,bold_prefix=None):
    para=doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        a=para.add_run(bold_prefix); set_font(a,bold=True); b=para.add_run(text[len(bold_prefix):]); set_font(b)
    else: para.add_run(text)
    return para

def bullets(items):
    for item in items: p(item,'List Bullet')

def numbers(items):
    numbering=doc.part.numbering_part.element
    ids=[int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    aid=max(ids or [0])+1
    abstract=OxmlElement('w:abstractNum'); abstract.set(qn('w:abstractNumId'),str(aid))
    multi=OxmlElement('w:multiLevelType'); multi.set(qn('w:val'),'singleLevel'); abstract.append(multi)
    lvl=OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'),'0'); abstract.append(lvl)
    start=OxmlElement('w:start'); start.set(qn('w:val'),'1'); lvl.append(start)
    fmt=OxmlElement('w:numFmt'); fmt.set(qn('w:val'),'decimal'); lvl.append(fmt)
    txt=OxmlElement('w:lvlText'); txt.set(qn('w:val'),'%1.'); lvl.append(txt)
    jc=OxmlElement('w:lvlJc'); jc.set(qn('w:val'),'left'); lvl.append(jc)
    pp=OxmlElement('w:pPr'); tabs=OxmlElement('w:tabs'); tab=OxmlElement('w:tab'); tab.set(qn('w:val'),'num'); tab.set(qn('w:pos'),'540'); tabs.append(tab); pp.append(tabs)
    ind=OxmlElement('w:ind'); ind.set(qn('w:left'),'540'); ind.set(qn('w:hanging'),'270'); pp.append(ind); lvl.append(pp); numbering.append(abstract)
    nums=[int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    nid=max(nums or [0])+1; num=OxmlElement('w:num'); num.set(qn('w:numId'),str(nid)); ref=OxmlElement('w:abstractNumId'); ref.set(qn('w:val'),str(aid)); num.append(ref); numbering.append(num)
    for item in items:
        para=p(item); para.paragraph_format.space_after=Pt(4); para.paragraph_format.line_spacing=1.18
        numPr=OxmlElement('w:numPr'); ilvl=OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0'); numid=OxmlElement('w:numId'); numid.set(qn('w:val'),str(nid)); numPr.append(ilvl); numPr.append(numid); para._p.get_or_add_pPr().append(numPr)

def table(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); trPr.append(repeat)
    for i,h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i],'1B211D'); q=t.rows[0].cells[i].paragraphs[0]; q.paragraph_format.space_after=Pt(0); rr=q.add_run(h); set_font(rr,size=9,bold=True,color=WHITE)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            if ridx%2: set_cell_shading(cells[i],'F7F4ED')
            q=cells[i].paragraphs[0]; q.paragraph_format.space_after=Pt(0); rr=q.add_run(str(val)); set_font(rr,size=9,color=NAVY)
    set_table_geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def callout(label,text):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; set_table_geometry(t,[9360]); c=t.cell(0,0); set_cell_shading(c,'F3E9D6'); q=c.paragraphs[0]; q.style='Callout'; a=q.add_run(label.upper()+'  '); set_font(a,size=9,bold=True,color=GOLD); b=q.add_run(text); set_font(b,size=10.5,color=NAVY); doc.add_paragraph().paragraph_format.space_after=Pt(1)

# Cover
p('MATCHDAY AFRICA','Title').alignment=WD_ALIGN_PARAGRAPH.CENTER
p('Complete App Feature, Workflow & Operations Guide','Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
p('Football intelligence · community · predictions · creator economy · Matchday War','Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')
callout('Purpose','A single reference for product owners, administrators, developers, content teams, support staff, partners, sponsors and future investors. It describes what the application does, how users experience it, how data moves through it, and what must be configured in production.')
table(['Document','Detail'],[
    ['Application','Matchday Africa with the Matchday War experience at /war'],
    ['Platform','Laravel 12, Blade, Vite, JavaScript, relational database'],
    ['Guide version',date.today().strftime('%d %B %Y')],
    ['Audience','Product, engineering, editorial, growth, commerce and operations'],
    ['Status','Features verified from the current local codebase'],
],[1800,7560])
doc.add_page_break()

p('How to use this guide','Heading 1')
p('This document is both a feature catalogue and an operating manual. Sections 1–5 explain the customer experience; sections 6–10 explain engagement, content, monetisation and administration; sections 11–15 cover architecture, integrations, scheduling, security and deployment.')
p('Contents at a glance','Heading 2')
table(['Section','Coverage'],[
['1–2','Product vision, audiences, navigation and account lifecycle'],['3–5','Football data, match centre, predictions and Matchday War'],['6–8','Community, personalization, discovery, creators and editorial'],['9–10','Growth engine, Premium, shop, sponsors and creator earnings'],['11–12','Admin controls, data model and system architecture'],['13–15','Automation, integrations, security, deployment and troubleshooting'],['Appendices','Myth Grammar, route map, operations checklist and glossary']], [1500,7860])
callout('Scope note','The guide describes implemented behaviour. A feature that depends on an external provider is identified as “configuration-dependent”; it becomes operational after valid credentials, callbacks, scheduler entries and/or webhooks are installed.')

p('1. Product overview','Heading 1')
p('Matchday Africa is a football platform designed to go beyond a conventional scores application. It combines live and scheduled match information with prediction games, supporter identity, social participation, African player discovery, editorial storytelling, a mythology-inspired game, and a commerce layer.')
p('Core product pillars','Heading 2')
table(['Pillar','What the app provides','Primary value'],[
['Match intelligence','Fixtures, live status, scores, events, previews, standings, teams and squads','A reliable matchday destination'],
['Participation','Predictions, leaderboards and private mini-leagues','Repeat visits and friendly competition'],
['Identity','Favourite clubs, supporter profiles, factions, points, streaks and badges','Belonging and retention'],
['Storytelling','Articles, match previews, creator publishing and Myth Grammar','A distinctive editorial voice'],
['Matchday War','Character selection, CPU play, browser two-player rooms and match-linked challenges','Shareable entertainment'],
['Growth','Referrals, email capture, campaign drafts, analytics and sharing','Acquisition and measurable distribution'],
['Commerce','Premium, rights-safe digital products, merchandise, sponsorship and creator revenue','Diversified monetisation'],
],[1450,4550,3360])
p('Primary audiences','Heading 2')
bullets(['Casual supporters who want fixtures and scores quickly.','Committed fans who want previews, context, standings, squad information and alerts.','Competitive users who make predictions and join private leagues.','Mobile-first users who play Matchday War and share challenges.','African football audiences discovering players across clubs and countries.','Creators who submit stories and earn from qualified store referrals.','Sponsors seeking measurable placements around matchday attention.','Administrators responsible for data quality, campaigns, editorial review and revenue operations.'])

p('2. Navigation, accounts and personalization','Heading 1')
p('The public navigation connects Home, Matches, Community/Leagues, War, African Players, Teams, Articles, Shop and Premium. Authenticated users additionally access predictions, their dashboard, private leagues, notification settings, profile and digital library.')
p('Authentication and account management','Heading 2')
bullets(['Email/password registration, login and logout.','Password reset, password confirmation and password update.','Email verification routes and resend workflow.','Google sign-in and account linking when Google OAuth credentials are configured.','Twitter/X OAuth authorization, callback and revocation when configured.','Profile editing and account deletion.','Role-aware access: ordinary users and administrators.'])
p('Personalized onboarding','Heading 2')
numbers(['A signed-in user opens onboarding.','The app lists active teams and allows the user to select between one and eight.','Selections are stored as polymorphic favourites.','The home experience can prioritize followed teams and the user can follow or unfollow a team from team pages.'])
p('Supporter identity','Heading 2')
bullets(['Public supporter pages show activity and prediction statistics.','Users can set a unique username, country, city, short biography and preferred War faction.','Daily Flame awards 10 Matchday Points once per day.','Consecutive daily claims increase current and best streaks.','A seven-day streak automatically grants the Seven-Day Watch badge.'])

p('3. Football data and the match centre','Heading 1')
p('The match centre is the application’s factual backbone. It stores leagues, teams, matches, events, standings and players locally, while synchronization services obtain current data from configured football providers.')
p('Home experience','Heading 2')
bullets(['Today’s fixtures with home and away teams, league, score and goal events.','Live-match emphasis for in-play statuses.','Featured leagues and editorial content.','Featured match previews.','Personalized sections for signed-in supporters.','Optional sponsored home placement with impression tracking.'])
p('Match list and detail','Heading 2')
bullets(['Date-based fixture browsing and live indicators.','Score, status, minute, teams and league context.','Goal, penalty, own-goal and disciplinary events when supplied.','Generated preview content and featured previews.','Public match chat viewing; authenticated message posting.','GIF search, trending GIFs and football GIFs through Giphy when configured.','Share links and social metadata for richer previews on social platforms.','Enhanced match routes for richer display variants.'])
p('Leagues, standings, teams and players','Heading 2')
bullets(['League directory, league detail and league standings.','Team directory, team profile and squad view.','Standing synchronization for configured leagues.','Player synchronization for team squads.','Country and team metadata stored locally for resilient page rendering.'])
callout('Data freshness','The scheduler refreshes today’s matches every 15 minutes and again hourly during peak hours. Standings update daily and player squads weekly. Production must run Laravel’s scheduler every minute for these timings to execute.')

p('4. Prediction system','Heading 1')
p('Predictions are organized into administrator-created prediction sets containing one or more matches. Users submit predictions before the deadline, the scoring service evaluates completed matches, and leaderboards aggregate earned points.')
p('Supported prediction types','Heading 2')
table(['Type','User predicts','Scoring behaviour'],[
['Result','Home win, draw or away win','Configured result points when correct.'],
['Exact score','A score such as 2–1','Configured exact-score points; one consolation point for the correct outcome when the score is wrong.'],
['First goalscorer','A named player','Scored only when sufficient event data exists; missing event data remains retryable.'],
['Total goals','A configured total-goal answer','Evaluated against the completed match total.'],
],[1500,2860,5000])
p('User prediction workflow','Heading 2')
numbers(['The user opens an active prediction set.','The app loads attached matches and any previous submission.','The user enters one or more allowed prediction types.','Server-side validation confirms the match exists, value is present and deadline has not passed.','The service creates or updates the predictions and reports submitted counts and validation errors.','After the match finishes, scheduled scoring awards points, updates leaderboards and sends notifications.'])
p('Competition and analytics','Heading 2')
bullets(['Global leaderboard with period and prediction-set filters.','Personal prediction history with date, correctness and set filters.','Personal statistics: total predictions, correct predictions, points, accuracy and rank.','Private or public prediction groups with a six-character invite code.','Group leaderboard based on members’ accumulated prediction points.','Administrator analytics, exports, match-level detail, transparency view and rescoring controls.','Deadline reminders and score-update notifications.'])
callout('Scoring integrity','The implemented scoring state prevents already-scored predictions from being counted again. Administrators can verify results, force scoring or rescore when provider data is corrected.')

p('5. Matchday War','Heading 1')
p('Matchday War translates football rivalry into a cinematic army-and-warrior experience. It lives natively at /war and uses the same Matchday database, fixtures, teams and growth analytics as the main site.')
p('Character and matchweek presentation','Heading 2')
bullets(['A roster of historical/fantasy warrior archetypes mapped to football teams in the current experience.','Card-style browsing optimized for mobile interaction.','Matchweek fixtures can present opposing warriors face to face.','Character imagery recedes into the battle background after kick-off while five tactical avatars per side represent the active formation.','Three attacking avatars and two defending avatars move in response to tactical actions.','Rights-safe wallpaper variants remove club names and protected badges for download and merchandise use.','A character trailer/video asset and atmospheric audio assets can be presented within the War experience.'])
p('Single-player game','Heading 2')
bullets(['Player selects a team/warrior and faces a computer opponent.','A football sits on the pitch and visibly moves toward either goal as possession and field position change.','ATTACK advances the ball when the player has possession.','TACKLE/DEFEND attempts to recover possession.','COUNTER attempts to win possession and immediately gain territory.','SHOOT becomes meaningful in the attacking zone and can produce a goal or a save.','The scoreboard uses team names rather than “You” and “CPU”.','A 60-second timer represents a full football match in compressed form.','Goal overlays, movement cues, highlighted controls and CPU animation make state changes visible.','Optional audio commentary is off by default and can be enabled in settings.'])
p('Browser-based two-player game','Heading 2')
numbers(['The host creates a room, selects a team and receives a six-character code.','The second player opens the site on another browser/device and joins using the code.','The host starts only after the guest is connected.','Both browsers poll the shared server room state.','Actions are turn-based and validated against player identity, possession and match state.','The server updates ball position, possession, score, commentary event, turn and remaining time.','Rooms expire after two hours and scheduled cleanup removes abandoned rooms.'])
callout('Two-player requirement','This is browser/server multiplayer, not Bluetooth. Both devices must reach the same deployed application and database. For local mobile testing they must be on the same network and use the computer’s LAN IP; for public play they use the live domain.')
p('Game rules implemented on the server','Heading 2')
table(['Action','Valid situation','Result'],[
['Attack','Side has possession','Ball advances 14 field units toward the opponent goal.'],['Defend/Tackle','Side does not have possession','72% recovery chance; successful tackle reclaims the ball.'],['Counter','Side does not have possession','58% recovery chance; success also gains 10 units of territory.'],['Shoot','Possession plus attacking zone','50% scoring chance; goal or save, then reset to halfway.'],['Invalid order','Wrong possession, range, turn or match state','Server rejects the action with an explanatory response.']], [1450,2900,5010])

p('6. Myth Grammar and storytelling','Heading 1')
p('Myth Grammar is the distinctive language layer that translates football events into cinematic, historically inspired beats. It is used for War narration, match stories and growth campaign copy while preserving real club/team names where clarity matters.')
table(['Football event','Mythic translation'],[
['Goal','A raid lands; warriors breach the line and steal plunder or a banner.'],['Assist','A flank manoeuvre draws the shield wall and opens a gap.'],['Red card','A warrior is captured and dragged from the field in chains.'],['Yellow card','A warrior is struck down, wounded, and rises again.'],['Penalty scored','A siege tower breaches the gate.'],['Penalty missed','The gate holds and the siege tower collapses.'],['Missed big chance','A spear sails wide of its mark.'],['Comeback','The retreating army rallies and turns the tide.'],['Draw','Both banners stand at dusk; the field is shared.'],['Clean sheet','The fortress wall holds unbreached.'],['Late winner','A final charge at sunset; the last arrow finds its mark.'],['Manager farewell','The old chieftain lays down his sword and passes the war-banner.'],['Trophy won','The war council crowns a high king and holds a feast.'],['Relegation','The army is banished from the great council to lesser lands.'],['Survival','The gates narrowly hold against winter.'],['Hat-trick','One warrior fells three enemies in a day.'],['Own goal','A warrior’s blade turns against a shield-brother by cruel accident.']], [2400,6960])
p('Current generated match story behaviour','Heading 2')
bullets(['Events are read in minute and sort order.','Goals, penalty goals, yellow cards and red cards become narrated beats.','Finished matches receive a winner or shared-field ending.','Unplayed matches receive a gathering-at-the-gates introduction.','Stories preserve the event minute, event type and generated text for timeline presentation.'])

p('7. Community, chat and social sharing','Heading 1')
bullets(['Public match conversations can be read without an account.','Authenticated users can post match chat messages.','User search supports mentions and participant discovery.','Giphy integration supports searchable and trending reaction GIFs.','Share routes generate platform-specific sharing destinations.','Share counts and popular-content endpoints help identify content that travels.','Social meta components provide richer link previews.','Analytics records page views and share-start events for anonymous and signed-in visitors.'])
p('Notifications','Heading 2')
bullets(['Users choose match alerts, prediction reminders and a weekly digest.','Digest day can be Monday, Friday or Sunday.','Finished-match and scored-prediction notifications are generated by application services.','The scheduler checks prediction reminders and score updates hourly.','A daily scheduled command processes configured Matchday digest delivery.'])

p('8. African player discovery and creator ecosystem','Heading 1')
p('African player discovery gives the platform a differentiated editorial and data focus. The creator system then lets approved contributors build stories around that football ecosystem.')
p('African player discovery','Heading 2')
bullets(['Dedicated African Players directory.','Filter by supported African nationality code.','Search players by name.','Pagination optimized for larger datasets.','Player detail pages include club/team relationship.','Non-African player records are excluded from discovery detail routes.'])
p('Creator lifecycle','Heading 2')
numbers(['A registered supporter submits display name, biography, speciality and optional social link.','The application creates a pending creator profile.','An administrator approves or rejects the application.','Approved creators access Creator Studio.','Creators submit title, excerpt and full story; each submission enters pending editorial review.','Administrators publish the story or return it to draft.','Published posts appear on creator profiles and the public article system.'])
p('Creator commerce','Heading 2')
bullets(['Each approved creator receives a shop URL containing their creator slug.','A purchase through that link associates the order with the creator.','After verified payment, the platform creates a commission record.','The default revenue share is 10%, configurable through CREATOR_REVENUE_SHARE.','Creators see lifetime and pending earnings in Creator Studio.','Administrators mark commission records paid after off-platform payout.'])

p('9. Growth engine','Heading 1')
p('The growth engine turns fixtures and War challenges into acquisition loops rather than treating every visit as isolated traffic.')
p('Acquisition and referral tracking','Heading 2')
bullets(['War email capture records consent and supports later audience activation.','Referral journeys have a UUID plus fixture, source and campaign attribution.','The funnel records landed, challenge, joined and completed timestamps.','Anonymous analytics uses a first-party visitor cookie; authenticated events also attach the user ID.','Popular-content and share-count endpoints surface high-performing pages.'])
p('Campaign generation','Heading 2')
numbers(['An administrator selects Generate in the War growth desk.','The system takes up to three high-priority fixtures within the next ten days.','It creates pre-match and post-match drafts for X, Instagram and TikTok.','Pre-match copy invites supporters to choose an army.','Post-match copy uses the result and Myth Grammar style.','An administrator approves or rejects each draft before external publication.'])
p('Automated distribution','Heading 2')
bullets(['Daily match-link tweeting at 08:00 when Twitter credentials are valid.','A 30-minute scheduled scan finds matches starting within two hours and prepares pre-match tweets.','Match pages and challenge URLs supply shareable destinations.','Campaign drafts remain human-approved; generation does not itself guarantee external publishing.'])

p('10. Monetisation and commerce','Heading 1')
p('Phase 3 introduces a shared Stripe-ready commerce system across subscriptions, digital goods, physical goods, sponsor inventory and creator commissions.')
p('Premium membership','Heading 2')
bullets(['Public Premium landing page with monthly offer and benefits.','Configured launch price: £3.99 per month.','Stripe subscription Checkout when STRIPE_PREMIUM_PRICE_ID is configured.','Verified payment stores Stripe customer/subscription IDs, active status and premium end date.','Subscription deletion webhook cancels Premium access.','Premium benefits are positioned as advanced intelligence, exclusive War drops and an ad-light experience.'])
p('Rights-safe shop','Heading 2')
table(['Product category','Current launch products','Fulfilment'],[
['Digital','Warrior Wallpaper Pack; Founders Digital Card Set','Verified payment grants a user entitlement and ZIP download in the Digital Vault.'],['Physical','War Council A3 Art Print; Matchday Legion Shirt','Stripe collects shipping address for supported countries; operational fulfilment is manual.'],['Membership','Matchday Africa Premium','Subscription activation updates the user’s premium status.']], [1600,3200,4560])
p('Payment and fulfilment flow','Heading 2')
numbers(['A signed-in user chooses a product or membership.','The app creates a pending local order and line item.','The server creates a Stripe Checkout Session and redirects the browser.','Stripe returns the user to the success page, but this redirect is not treated as proof of payment.','Stripe posts a signed webhook to /stripe/webhook.','The app verifies the signature and accepts completed or asynchronous payment success events.','CommerceService marks the order paid and performs idempotent fulfilment.','Digital purchases grant one entitlement; memberships activate Premium; creator-attributed orders create one commission.'])
callout('Rights-safe policy','Commercial products must use original Matchday Africa artwork without club names, club crests, league marks or other protected identity. Existing rights-safe artwork is stored under public/war/downloads/rights-safe.')
p('Sponsor inventory','Heading 2')
bullets(['Admin-created placements support home, predictions, matches and War slots.','Each placement stores sponsor, headline, destination, optional image, start/end dates and active state.','A live slot selects an eligible campaign and records an impression.','Clicks pass through a tracked redirect and increment the campaign click count.','Sponsor links are marked sponsored/nofollow in rendered markup.'])

p('11. Administration and operational controls','Heading 1')
table(['Admin area','Capabilities'],[
['Dashboard/System','Application overview, API status and cache clearing.'],['Data sync','Manual league, match, standing and player synchronization.'],['Matches','Browse, update scores, bulk-update statuses, verify, force-score and auto-update.'],['Predictions','Create/edit sets, attach matches, activate/close/archive, score/rescore, analytics, transparency and export.'],['Match previews','Generate daily, regenerate, feature/unfeature, inspect statistics and remove.'],['Editorial','Create/edit/publish articles, preview, moderate creator applications and submissions.'],['War','Review War configuration and growth campaign operations.'],['Twitter','Authorize, test connection, send test posts and trigger match posts.'],['Commerce','Revenue summary, orders, catalogue availability, sponsor creation and creator payout status.']], [1700,7660])
p('Recommended daily operator routine','Heading 2')
numbers(['Check API status and today’s fixture ingestion.','Review live-match status anomalies and missing events.','Confirm active prediction sets and deadlines.','Moderate creator submissions and match preview quality.','Approve relevant War campaign drafts.','Review failed/pending Stripe orders and new creator liabilities.','Check sponsor pacing, impressions and clicks.','Respond to user reports, then verify logs and scheduled-task health.'])

p('12. Technical architecture','Heading 1')
p('The application follows Laravel’s MVC/service architecture. Routes receive requests, controllers validate and coordinate, services contain integration and business logic, Eloquent models persist state, Blade renders pages, and Vite builds CSS/JavaScript assets.')
table(['Layer','Implemented responsibility'],[
['Presentation','Blade templates, reusable components, responsive Matchday/War theme, Vite-built CSS and JavaScript.'],['HTTP','Public, authenticated, API and admin routes; controllers; CSRF protection; admin middleware.'],['Domain services','Football synchronization, previews, predictions, analytics, Myth Grammar, social sharing, Stripe checkout and commerce fulfilment.'],['Persistence','Relational tables for football, community, prediction, War, growth, content and commerce domains.'],['Automation','Laravel scheduler, Artisan commands, queued preview generation and event/notification dispatch.'],['External boundary','Football Data, Odds API, OpenRouter, Giphy, Google OAuth, Twitter/X, Stripe and email services.']], [1800,7560])
p('Major data domains','Heading 2')
bullets(['Football: leagues, teams, matches, match events, standings and players.','Prediction: sets, set matches, user predictions, results and leaderboards.','Community: users, favourites, chats, groups, badges and notification preferences.','Editorial: blogs, match previews, match stories and creator profiles.','War/growth: factions, rooms, campaigns, referrals and subscribers.','Commerce: products, orders, order items, entitlements, sponsors and creator earnings.','Measurement: social shares and analytics events.'])
p('Key state transitions','Heading 2')
table(['Domain','Lifecycle'],[
['Match','Scheduled → live/in play → finished; score/events may update repeatedly.'],['Prediction set','Draft → active → closed → archived; individual predictions become scored after match completion.'],['Creator','Pending → approved/rejected; submission draft → pending review → published/returned.'],['War room','Waiting → ready → playing → finished/expired.'],['Order','Pending → paid; failed/refunded states are reserved in the data model.'],['Sponsor','Inactive/scheduled → live during date window → ended.']], [1800,7560])

p('13. Integrations and configuration','Heading 1')
table(['Integration','Purpose','Important environment values'],[
['Football Data','Leagues, fixtures, scores, standings, teams and players','FOOTBALL_DATA_API_KEY, base URL, cache/rate settings'],['Odds API','Weekend, upcoming and match odds endpoints','ODDS_API_KEY, URL, regions, markets'],['OpenRouter','AI-assisted match preview generation','OPENROUTER_API_KEY, model, daily limit'],['Giphy','Chat reactions','GIPHY_API_KEY'],['Google OAuth','Authentication/account linking','GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, redirect URI'],['Twitter/X','OAuth and automated distribution','Twitter client/API credentials and tokens'],['Stripe','Checkout, subscription and verified fulfilment','STRIPE_KEY, STRIPE_SECRET, STRIPE_WEBHOOK_SECRET, STRIPE_PREMIUM_PRICE_ID'],['Email','Password resets, notifications and digest','Configured Laravel mail transport/Resend/Postmark/SES/Mailgun values']], [1550,3200,4610])
callout('Secrets','Production secrets belong only in the server environment. Do not commit the live .env file. After changing environment values, clear and rebuild Laravel configuration cache.')

p('14. Scheduling, queues and data freshness','Heading 1')
table(['Frequency','Task'],[
['Every 5 minutes','Score eligible predictions.'],['Every 15 minutes','Synchronize today’s matches.'],['Every 30 minutes','Look for matches within two hours and trigger pre-match Twitter workflow.'],['Hourly','Send prediction reminders/score notifications; clean expired War rooms.'],['14:00–23:00 hourly','Additional peak-hours match synchronization.'],['00:00 daily','Clear application cache.'],['03:00 daily','Synchronize standings.'],['08:00 daily','Process Matchday digest and scheduled match-link tweets.'],['Sunday 04:00','Synchronize player squads.'],['Weekly','Clean old application logs.']], [2150,7210])
p('Required server jobs','Heading 2')
bullets(['Cron every minute: php artisan schedule:run.','A long-running queue worker when queued preview generation or asynchronous jobs are used: php artisan queue:work --tries=3.','Log rotation/monitoring appropriate for cPanel or the production host.','HTTPS and a public callback URL for OAuth and Stripe webhooks.'])

p('15. Security, privacy and reliability','Heading 1')
bullets(['Laravel CSRF protection covers browser write routes; only the Stripe webhook path is excluded because Stripe authenticates with a signature.','Passwords are hashed through Laravel authentication.','Admin routes require authentication plus AdminMiddleware.','Two-player War actions validate a UUID player token, room membership, turn and state.','Checkout requires authentication, and digital downloads require a matching entitlement.','Stripe webhook signatures are verified before fulfilment.','Commerce fulfilment is idempotent to prevent duplicate grants or commissions from repeat webhooks.','Queries validate model existence and input constraints.','Analytics stores a first-party visitor identifier and optional user ID; production privacy notices and consent controls should describe this.','War subscriber records include consent and unsubscribe timestamps.','Scheduled tasks log errors so operational failures can be diagnosed.'])
p('Known operational boundaries','Heading 2')
bullets(['Physical merchandise fulfilment is not connected to a print-on-demand provider; staff must fulfil paid orders manually or add a future integration.','Creator payout marking is an internal ledger action; it does not transfer money automatically.','Premium currently records a one-month access horizon when checkout completes; advanced renewal period synchronization can be expanded with invoice/subscription update webhooks.','The two-player game uses HTTP room polling and a shared database; it is not peer-to-peer or Bluetooth.','External-data accuracy depends on provider coverage, quota and event completeness.','Audio commentary quality depends on browser speech/audio support and the assets selected for production.'])

p('16. Deployment to cPanel / matchday.africa','Heading 1')
p('The Laravel application is intended to run from the domain root, with Matchday War available at /war. The web server document root must point to Laravel’s public directory, never the repository root.')
p('Deployment sequence','Heading 2')
numbers(['Upload or deploy the application outside the public web root where practical.','Point the domain document root to matchday-africa/public.','Create the production database and set DB_CONNECTION plus credentials in .env.','Set APP_ENV=production, APP_DEBUG=false, APP_URL=https://matchday.africa and generate APP_KEY.','Install PHP dependencies with Composer’s optimized production flags.','Install/build frontend dependencies or upload the compiled public/build directory.','Run all database migrations with --force.','Create storage link, set writable permissions for storage and bootstrap/cache.','Configure all required provider credentials and exact OAuth callback URLs.','Register the Stripe webhook at https://matchday.africa/stripe/webhook.','Create scheduler and queue-worker jobs.','Cache configuration, routes and views, then run health and feature checks.'])
p('Essential commands','Heading 2')
for cmd in ['composer install --no-dev --optimize-autoloader','php artisan key:generate  (first deployment only)','php artisan migrate --force','php artisan storage:link','npm ci && npm run build','php artisan config:cache','php artisan route:cache','php artisan view:cache','php artisan test']:
    q=p(cmd); q.style='Callout'
p('Post-deployment acceptance checks','Heading 2')
bullets(['Home, login, registration, matches, leagues, teams, articles, players, shop, Premium and /war return successfully.','A test account can follow a team, predict, claim Daily Flame and join a private prediction league.','Administrator can sync data and access all admin pages.','Stripe test checkout reaches the success page and a signed webhook changes the order to paid.','Purchased digital item appears in the user library and downloads.','Two browsers can create, join and complete a War room.','Scheduler and queue logs show successful execution.','Mobile layouts remain usable without horizontal overflow.'])

p('Appendix A. Public route guide','Heading 1')
table(['Area','Main routes'],[
['Core','/, /matches, /matches/{id}, /leagues, /leagues/{id}, /teams, /teams/{id}'],['Content','/blogs, /blogs/{slug}, /african-players, /african-players/{id}, /creators/{slug}'],['Participation','/predictions, /predictions/history, /predictions/leaderboard, /prediction-leagues'],['Identity','/onboarding, /supporters/{id}, /profile, /notification-settings, /dashboard'],['War','/war, /war/match/{id}, /war/challenge/{id}, /war/api/fixtures, /war/api/rooms'],['Commerce','/premium, /shop, /checkout/{product}, /library, /stripe/webhook'],['Social','/share/{type}/{id}/{platform}, /matches/{id}/chats, /gifs/*']], [1750,7610])
p('Appendix B. Operations checklist','Heading 1')
p('Before every matchweek','Heading 2')
bullets(['Confirm fixtures are synchronized through the next 10–14 days.','Verify featured leagues, team mappings and War factions.','Create/activate prediction sets and check deadlines.','Generate previews and inspect factual accuracy.','Generate War campaigns, revise and approve selected drafts.','Confirm sponsor date windows and destination URLs.','Check scheduler, queue, mail and provider quotas.'])
p('After matches finish','Heading 2')
bullets(['Verify final scores and important match events.','Confirm prediction scoring completed and leaderboards moved.','Resolve unscored first-goalscorer predictions when event data was delayed.','Review generated Myth Grammar stories and post-match campaigns.','Check War referral completion and sharing performance.','Review commerce, sponsor and creator-revenue dashboards.'])
p('Appendix C. Glossary','Heading 1')
table(['Term','Meaning'],[
['Prediction set','An administrator-defined collection of matches and prediction opportunities.'],['War faction','The warrior identity associated with a football team.'],['Myth Grammar','The translation rules turning football events into mythic narrative beats.'],['War room','A temporary server-side state container for a browser two-player match.'],['Entitlement','A record proving a user owns a digital product and may download it.'],['Sponsor slot','A named page location that can display one eligible sponsor placement.'],['Creator earning','A commission ledger entry created by a paid creator-attributed order.'],['Idempotent fulfilment','Processing that safely produces the same result when a webhook is delivered more than once.']], [2000,7360])
p('Document control','Heading 1')
p('This guide was generated from the application routes, controllers, services, models, migrations, views and scheduler present in the local Matchday Africa codebase. It should be updated whenever a phase changes user-facing behaviour, external integrations, database structure, scheduled tasks or deployment requirements.')

doc.core_properties.title='Matchday Africa — Complete App Feature, Workflow & Operations Guide'
doc.core_properties.subject='Comprehensive product and technical documentation'
doc.core_properties.author='Matchday Africa'
doc.core_properties.keywords='Matchday Africa, football, predictions, Matchday War, Stripe, operations'
doc.save(OUT)
print(OUT)
